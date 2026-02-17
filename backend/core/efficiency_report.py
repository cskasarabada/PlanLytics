# core/efficiency_report.py
"""
Configuration Efficiency Report — analyzes an oracle_mapping dict and produces
actionable recommendations for optimizing Oracle ICM plan configurations.

Categories of analysis:
  1. Duplicate / redundant expressions
  2. Unused objects (rate tables, expressions, measures not referenced)
  3. Simplification opportunities (flat vs tiered, expression consolidation)
  4. Credit category rationalization
  5. Calculation settings optimization
  6. Missing best-practice patterns
"""
from typing import Dict, Any, List
from collections import Counter, defaultdict
from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ---------------------------------------------------------------------------
# Core analysis engine
# ---------------------------------------------------------------------------

def analyze_efficiency(oracle_mapping: dict) -> Dict[str, Any]:
    """Analyze oracle_mapping and return structured efficiency findings.

    Returns dict with:
        - score: 0-100 overall efficiency score
        - findings: list of dicts with category, severity, title, detail, recommendation
        - summary: high-level stats
    """
    findings: List[Dict[str, str]] = []
    om = oracle_mapping

    plans = om.get("compensation_plans", [])
    components = om.get("plan_components", [])
    measures = om.get("performance_measures", [])
    expressions = om.get("expressions", [])
    rate_tables = om.get("rate_tables", [])
    rate_table_rates = om.get("rate_table_rates", [])
    rate_dims = om.get("rate_dimensions", [])
    credit_cats = om.get("credit_categories", [])
    calc_settings = om.get("calculation_settings", [])
    goals = om.get("performance_goals", [])
    scorecards = om.get("scorecards", [])

    # ── 1. Duplicate expression detection ──────────────────────
    _check_duplicate_expressions(expressions, findings)

    # ── 2. Unused object detection ──────────────────────────────
    _check_unused_rate_tables(rate_tables, components, findings)
    _check_unused_expressions(expressions, components, measures, findings)
    _check_unused_measures(measures, components, findings)

    # ── 3. Simplification opportunities ─────────────────────────
    _check_single_tier_rate_tables(rate_tables, rate_dims, rate_table_rates, findings)
    _check_flat_calculation_candidates(components, rate_tables, findings)
    _check_expression_consolidation(expressions, findings)

    # ── 4. Credit category rationalization ──────────────────────
    _check_credit_category_usage(credit_cats, measures, findings)

    # ── 5. Calculation settings optimization ────────────────────
    _check_calculation_settings(calc_settings, components, findings)
    _check_missing_true_up(components, findings)
    _check_indirect_credit_consistency(components, findings)

    # ── 6. Missing best-practice patterns ───────────────────────
    _check_missing_goals(measures, goals, findings)
    _check_missing_scorecards(measures, scorecards, findings)
    _check_component_naming(components, findings)
    _check_expression_categories(expressions, findings)

    # ── Score calculation ───────────────────────────────────────
    # Start at 100, deduct based on severity
    score = 100
    for f in findings:
        sev = f.get("severity", "info")
        if sev == "high":
            score -= 8
        elif sev == "medium":
            score -= 4
        elif sev == "low":
            score -= 2
        else:
            score -= 1
    score = max(0, min(100, score))

    return {
        "score": score,
        "findings": findings,
        "summary": {
            "total_findings": len(findings),
            "high_severity": sum(1 for f in findings if f["severity"] == "high"),
            "medium_severity": sum(1 for f in findings if f["severity"] == "medium"),
            "low_severity": sum(1 for f in findings if f["severity"] == "low"),
            "info": sum(1 for f in findings if f["severity"] == "info"),
            "total_objects": (
                len(plans) + len(components) + len(measures) +
                len(expressions) + len(rate_tables) + len(rate_table_rates) +
                len(rate_dims) + len(credit_cats) + len(calc_settings) +
                len(goals) + len(scorecards)
            ),
            "compensation_plans": len(plans),
            "plan_components": len(components),
            "performance_measures": len(measures),
            "expressions": len(expressions),
            "rate_tables": len(rate_tables),
            "credit_categories": len(credit_cats),
        },
    }


# ---------------------------------------------------------------------------
# Check functions
# ---------------------------------------------------------------------------

def _check_duplicate_expressions(expressions: list, findings: list):
    """Find expressions that have identical formulas (same detail type + attributes)."""
    # Group expression detail rows by expression name
    by_name: Dict[str, list] = defaultdict(list)
    for e in expressions:
        name = e.get("expression_name", "")
        by_name[name].append(e)

    # Build a signature per expression name
    signatures: Dict[str, str] = {}
    for name, rows in by_name.items():
        sig_parts = []
        for r in sorted(rows, key=lambda x: x.get("sequence", 0)):
            sig_parts.append(
                f"{r.get('expression_detail_type', '')}|"
                f"{r.get('basic_attributes_group', '')}|"
                f"{r.get('basic_attribute_name', '')}|"
                f"{r.get('measure_name', '')}|"
                f"{r.get('expression_operator', '')}|"
                f"{r.get('constant_value', '')}"
            )
        signatures[name] = "::".join(sig_parts)

    # Find duplicates
    sig_to_names: Dict[str, List[str]] = defaultdict(list)
    for name, sig in signatures.items():
        sig_to_names[sig].append(name)

    for sig, names in sig_to_names.items():
        if len(names) > 1:
            findings.append({
                "category": "Duplicate Expressions",
                "severity": "medium",
                "title": f"{len(names)} expressions have identical formulas",
                "detail": f"Expressions {', '.join(names)} compute the same formula. "
                          f"Consider consolidating into a single shared expression.",
                "recommendation": f"Merge into one expression and reference it from all components.",
            })


def _check_unused_rate_tables(rate_tables: list, components: list, findings: list):
    """Find rate tables not referenced by any plan component."""
    referenced = {pc.get("rate_table_name", "") for pc in components} - {""}
    for rt in rate_tables:
        name = rt.get("rate_table_name", rt.get("table_name", ""))
        if name and name not in referenced:
            findings.append({
                "category": "Unused Objects",
                "severity": "medium",
                "title": f"Rate table '{name}' is not used by any component",
                "detail": f"Rate table '{name}' exists but no plan component references it.",
                "recommendation": "Remove unused rate table to reduce deployment complexity.",
            })


def _check_unused_expressions(
    expressions: list, components: list, measures: list, findings: list,
):
    """Find expressions not referenced by any component or measure."""
    # Collect all expression names
    expr_names = {e.get("expression_name", "") for e in expressions} - {""}

    # Collect referenced expressions
    referenced = set()
    for pc in components:
        ref = pc.get("incentive_formula_expression", "")
        if ref:
            referenced.add(ref)
        ref = pc.get("rate_dimension_input_expression")
        if ref:
            referenced.add(ref)
    for pm in measures:
        ref = pm.get("measure_formula_expression_name", "")
        if ref:
            referenced.add(ref)

    unused = expr_names - referenced
    if unused:
        findings.append({
            "category": "Unused Objects",
            "severity": "low",
            "title": f"{len(unused)} expression(s) not directly referenced",
            "detail": f"Expressions: {', '.join(sorted(unused))}. These may be intermediate "
                      f"expressions referenced within other formulas, or truly unused.",
            "recommendation": "Verify these are needed; remove orphaned expressions.",
        })


def _check_unused_measures(measures: list, components: list, findings: list):
    """Find performance measures not referenced by any plan component."""
    referenced = {pc.get("performance_measure_name", "") for pc in components} - {""}
    for pm in measures:
        name = pm.get("name", "")
        if name and name not in referenced:
            findings.append({
                "category": "Unused Objects",
                "severity": "medium",
                "title": f"Performance measure '{name}' is not used by any component",
                "detail": f"Measure '{name}' exists but no plan component references it.",
                "recommendation": "Remove unused measure or link it to a component.",
            })


def _check_single_tier_rate_tables(
    rate_tables: list, rate_dims: list, rate_table_rates: list, findings: list,
):
    """Flag rate tables with only 1 tier — these could be flat rates instead."""
    # Count tiers per rate table
    tiers_per_table: Counter = Counter()
    for rtr in rate_table_rates:
        tiers_per_table[rtr.get("rate_table_name", "")] += 1

    for rt in rate_tables:
        name = rt.get("rate_table_name", rt.get("table_name", ""))
        if name and tiers_per_table.get(name, 0) <= 1:
            findings.append({
                "category": "Simplification",
                "severity": "low",
                "title": f"Rate table '{name}' has only 1 tier",
                "detail": f"A single-tier rate table is functionally a flat rate. "
                          f"Consider using a flat calculation method instead.",
                "recommendation": "Replace single-tier rate table with a flat rate to simplify config.",
            })


def _check_flat_calculation_candidates(components: list, rate_tables: list, findings: list):
    """Components using Tiered method but with no rate table."""
    for pc in components:
        method = pc.get("calculation_method", "")
        rt = pc.get("rate_table_name", "")
        if method == "Tiered" and not rt:
            findings.append({
                "category": "Simplification",
                "severity": "low",
                "title": f"Component '{pc.get('plan_component_name', '')}' is Tiered without a rate table",
                "detail": "Calculation method is 'Tiered' but no rate table is assigned.",
                "recommendation": "Either assign a rate table or switch to 'Flat' calculation method.",
            })


def _check_expression_consolidation(expressions: list, findings: list):
    """Find expressions that could be consolidated (e.g., many SUM(Credit.Credit Amount))."""
    # Group by formula pattern
    pattern_count: Counter = Counter()
    pattern_names: Dict[str, List[str]] = defaultdict(list)

    by_name: Dict[str, list] = defaultdict(list)
    for e in expressions:
        by_name[e.get("expression_name", "")].append(e)

    for name, rows in by_name.items():
        if len(rows) == 1:
            r = rows[0]
            pattern = f"{r.get('expression_detail_type', '')}:{r.get('basic_attributes_group', '')}:{r.get('basic_attribute_name', '')}"
            if pattern != "::":
                pattern_count[pattern] += 1
                pattern_names[pattern].append(name)

    for pattern, count in pattern_count.items():
        if count >= 3:
            names = pattern_names[pattern][:5]
            findings.append({
                "category": "Simplification",
                "severity": "info",
                "title": f"{count} expressions use the same single-step pattern",
                "detail": f"Expressions like {', '.join(names)} all use pattern '{pattern}'. "
                          f"If they compute the same value, consider sharing one expression.",
                "recommendation": "Consolidate identical single-step expressions into a shared formula.",
            })


def _check_credit_category_usage(credit_cats: list, measures: list, findings: list):
    """Check for credit categories not linked to any measure."""
    used_cats = {m.get("credit_category_name", "") for m in measures} - {""}
    for cc in credit_cats:
        name = cc.get("credit_category_name", cc.get("name", ""))
        if name and name not in used_cats:
            findings.append({
                "category": "Credit Categories",
                "severity": "info",
                "title": f"Credit category '{name}' not linked to any measure",
                "detail": f"Category '{name}' exists but no performance measure references it.",
                "recommendation": "Verify this category is needed for transaction routing; "
                                  "remove if unused to reduce deployment scope.",
            })


def _check_calculation_settings(calc_settings: list, components: list, findings: list):
    """Check for missing or inconsistent calculation settings."""
    cs_names = {cs.get("plan_component_name", "") for cs in calc_settings}
    pc_names = {pc.get("plan_component_name", "") for pc in components}
    missing = pc_names - cs_names - {""}
    if missing:
        findings.append({
            "category": "Calculation Settings",
            "severity": "medium",
            "title": f"{len(missing)} component(s) missing calculation settings",
            "detail": f"Components without explicit settings: {', '.join(sorted(missing))}. "
                      f"Defaults will be used but explicit settings are recommended.",
            "recommendation": "Add calculation settings for all components to ensure correct behavior.",
        })


def _check_missing_true_up(components: list, findings: list):
    """Flag components that might benefit from true-up based on interval settings."""
    for pc in components:
        calc_incentive = pc.get("calculate_incentive", "Per interval")
        true_up = pc.get("true_up", "No")
        if calc_incentive == "Per interval" and true_up == "No":
            # Many orgs want true-up for period-based calculations
            pass  # Only flag if pattern is widespread
    # Count how many use Per interval without true-up
    no_trueup_count = sum(
        1 for pc in components
        if pc.get("calculate_incentive", "Per interval") == "Per interval"
        and pc.get("true_up", "No") == "No"
    )
    if no_trueup_count > 3:
        findings.append({
            "category": "Best Practice",
            "severity": "info",
            "title": f"{no_trueup_count} components use 'Per interval' without true-up",
            "detail": "Period-based calculations without true-up may lead to over/underpayment "
                      "when actuals are revised. Consider enabling true-up for critical components.",
            "recommendation": "Evaluate whether true-up should be enabled for payout accuracy.",
        })


def _check_indirect_credit_consistency(components: list, findings: list):
    """Check if indirect credit settings are mixed across components."""
    indirect_values = Counter(
        pc.get("include_indirect_credits", "None") for pc in components
    )
    if len(indirect_values) > 1 and "None" in indirect_values:
        non_none = {k: v for k, v in indirect_values.items() if k != "None"}
        if non_none:
            findings.append({
                "category": "Best Practice",
                "severity": "info",
                "title": "Mixed indirect credit settings across components",
                "detail": f"Some components use indirect credits ({dict(non_none)}) while others don't. "
                          f"Ensure this is intentional for your overlay/split crediting model.",
                "recommendation": "Verify indirect credit settings align with territory and overlay design.",
            })


def _check_missing_goals(measures: list, goals: list, findings: list):
    """Flag measures without performance goals."""
    measures_with_goals = {g.get("performance_measure_name", "") for g in goals}
    missing = []
    for pm in measures:
        name = pm.get("name", "")
        if name and name not in measures_with_goals:
            missing.append(name)
    if missing:
        findings.append({
            "category": "Missing Configuration",
            "severity": "high",
            "title": f"{len(missing)} measure(s) have no performance goals",
            "detail": f"Measures without goals: {', '.join(missing[:10])}. "
                      f"Without goals, attainment cannot be calculated properly.",
            "recommendation": "Add performance goals for all measures that require target-based attainment.",
        })


def _check_missing_scorecards(measures: list, scorecards: list, findings: list):
    """Check if scorecard patterns are partially configured."""
    if scorecards and len(scorecards) > 0:
        sc_measures = {sc.get("performance_measure_name", "") for sc in scorecards}
        measures_without_sc = [
            m.get("name", "") for m in measures
            if m.get("name", "") and m.get("name", "") not in sc_measures
        ]
        if measures_without_sc and len(measures_without_sc) < len(measures):
            findings.append({
                "category": "Missing Configuration",
                "severity": "low",
                "title": f"Scorecards configured for some but not all measures",
                "detail": f"Measures without scorecards: {', '.join(measures_without_sc[:5])}.",
                "recommendation": "Either extend scorecards to all measures or document which are excluded.",
            })


def _check_component_naming(components: list, findings: list):
    """Check for naming consistency issues."""
    names = [pc.get("plan_component_name", "") for pc in components]
    # Check for very long names
    long_names = [n for n in names if len(n) > 80]
    if long_names:
        findings.append({
            "category": "Best Practice",
            "severity": "low",
            "title": f"{len(long_names)} component(s) have names longer than 80 characters",
            "detail": f"Long names may cause display issues in Oracle ICM. "
                      f"Example: '{long_names[0][:60]}...'",
            "recommendation": "Shorten component names while keeping them descriptive.",
        })


def _check_expression_categories(expressions: list, findings: list):
    """Check for expressions with unexpected category assignments."""
    by_name: Dict[str, list] = defaultdict(list)
    for e in expressions:
        by_name[e.get("expression_name", "")].append(e)

    for name, rows in by_name.items():
        categories = set(r.get("expression_category", "") for r in rows)
        if len(categories) > 1:
            findings.append({
                "category": "Data Quality",
                "severity": "medium",
                "title": f"Expression '{name}' has mixed categories",
                "detail": f"Categories found: {', '.join(categories)}. "
                          f"All detail rows should have the same category.",
                "recommendation": "Align all expression detail rows to a single category.",
            })


# ---------------------------------------------------------------------------
# Document generation
# ---------------------------------------------------------------------------

def generate_efficiency_report_doc(
    report: Dict[str, Any],
    analysis: dict,
    analysis_id: str,
) -> bytes:
    """Generate a Word document (.docx) from the efficiency report."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    plan_name = _get_plan_name(analysis)

    # Title page
    for _ in range(4):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Configuration Efficiency Report")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(15, 23, 42)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(plan_name)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 116, 139)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    score = report.get("score", 0)
    color = RGBColor(34, 197, 94) if score >= 80 else (
        RGBColor(234, 179, 8) if score >= 60 else RGBColor(239, 68, 68)
    )
    run = p.add_run(f"Efficiency Score: {score}/100")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = color

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Analysis ID: {analysis_id}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(148, 163, 184)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(148, 163, 184)

    doc.add_page_break()

    # ── Executive Summary ──
    doc.add_heading("Executive Summary", level=1)
    summary = report.get("summary", {})
    doc.add_paragraph(
        f"This report analyzes the efficiency of the {plan_name} configuration "
        f"across {summary.get('total_objects', 0)} total ICM objects. "
        f"The overall efficiency score is {score}/100."
    )

    _add_table(doc, ["Metric", "Value"], [
        ["Total ICM Objects", str(summary.get("total_objects", 0))],
        ["Compensation Plans", str(summary.get("compensation_plans", 0))],
        ["Plan Components", str(summary.get("plan_components", 0))],
        ["Performance Measures", str(summary.get("performance_measures", 0))],
        ["Expressions", str(summary.get("expressions", 0))],
        ["Rate Tables", str(summary.get("rate_tables", 0))],
        ["Credit Categories", str(summary.get("credit_categories", 0))],
        ["", ""],
        ["Total Findings", str(summary.get("total_findings", 0))],
        ["High Severity", str(summary.get("high_severity", 0))],
        ["Medium Severity", str(summary.get("medium_severity", 0))],
        ["Low Severity", str(summary.get("low_severity", 0))],
        ["Informational", str(summary.get("info", 0))],
    ])
    doc.add_paragraph("")

    # ── Findings by Category ──
    doc.add_heading("Findings & Recommendations", level=1)
    findings = report.get("findings", [])
    if not findings:
        p = doc.add_paragraph("No efficiency issues found. Configuration looks well-optimized.")
        run = p.runs[0]
        run.font.color.rgb = RGBColor(34, 197, 94)
    else:
        # Group by category
        by_category: Dict[str, list] = defaultdict(list)
        for f in findings:
            by_category[f.get("category", "Other")].append(f)

        for category, cat_findings in by_category.items():
            doc.add_heading(category, level=2)
            for f in cat_findings:
                severity = f.get("severity", "info").upper()
                sev_color = (
                    RGBColor(239, 68, 68) if severity == "HIGH" else
                    RGBColor(234, 179, 8) if severity == "MEDIUM" else
                    RGBColor(100, 116, 139)
                )

                p = doc.add_paragraph()
                run = p.add_run(f"[{severity}] ")
                run.bold = True
                run.font.color.rgb = sev_color
                run = p.add_run(f.get("title", ""))
                run.bold = True

                doc.add_paragraph(f.get("detail", ""))

                p = doc.add_paragraph()
                run = p.add_run("Recommendation: ")
                run.bold = True
                run.font.color.rgb = RGBColor(14, 165, 233)
                p.add_run(f.get("recommendation", ""))
                doc.add_paragraph("")

    # ── Score Breakdown ──
    doc.add_heading("Score Methodology", level=1)
    doc.add_paragraph(
        "The efficiency score starts at 100 and deducts points based on findings: "
        "high severity (-8), medium (-4), low (-2), informational (-1). "
        "A score of 80+ indicates a well-optimized configuration."
    )

    # Footer
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Generated by ICM PlanLytics AI Plan Analytics")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(148, 163, 184)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _add_table(doc: Document, headers: List[str], rows: List[List[str]]):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value) if value is not None else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def _get_plan_name(analysis: dict) -> str:
    om = analysis.get("oracle_mapping", {})
    plans = om.get("compensation_plans", [])
    if plans:
        return plans[0].get("name", "Compensation Plan")
    return "Compensation Plan Analysis"
