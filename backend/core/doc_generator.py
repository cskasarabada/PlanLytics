# core/doc_generator.py
"""
Generate Design Document and Configuration Document from AI analysis output.

Design Document — summarizes plan mechanics, risks, stakeholder questions,
and recommendations for business review.

Configuration Document — detailed Oracle ICM object mapping with deployment
instructions, expression formulas, rate tables, and calculation settings.
"""
from io import BytesIO
from datetime import datetime
from typing import Dict, Any, List

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ─── Shared helpers ──────────────────────────────────────────────────

def _add_title_page(doc: Document, title: str, subtitle: str, analysis_id: str):
    """Add a formatted title page."""
    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(15, 23, 42)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Analysis ID: {analysis_id}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(148, 163, 184)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(148, 163, 184)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ICM PlanLytics — AI Plan Analytics")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(14, 165, 233)

    doc.add_page_break()


def _add_table(doc: Document, headers: List[str], rows: List[List[str]]):
    """Add a formatted table with headers and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value) if value is not None else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def _severity_label(severity: str) -> str:
    s = (severity or "").lower()
    if s in ("high", "critical"):
        return "HIGH"
    if s in ("medium", "moderate"):
        return "MEDIUM"
    return "LOW"


# ─── Design Document ────────────────────────────────────────────────

def generate_design_document(analysis: dict, analysis_id: str) -> bytes:
    """
    Generate a Design Document (.docx) from analysis results.

    Sections:
      1. Executive Summary
      2. Plan Structure & Mechanics
      3. Risk Assessment
      4. Stakeholder Questions
      5. Oracle ICM Object Summary
      6. Recommendations
      7. Data Integration Requirements
      8. Governance Controls
    """
    doc = Document()

    # Style defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    plan_name = _get_plan_name(analysis)
    _add_title_page(doc, "Design Document", plan_name, analysis_id)

    # ── Table of Contents placeholder ──
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Plan Structure & Mechanics",
        "3. Risk Assessment",
        "4. Stakeholder Questions",
        "5. Oracle ICM Object Summary",
        "6. Recommendations",
        "7. Data Integration Requirements",
        "8. Governance & Controls",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # ── 1. Executive Summary ──
    doc.add_heading("1. Executive Summary", level=1)
    om = analysis.get("oracle_mapping", {})
    n_plans = len(om.get("compensation_plans", []))
    n_components = len(om.get("plan_components", []))
    n_measures = len(om.get("performance_measures", []))
    n_expressions = len(om.get("expressions", []))
    n_rate_tables = len(om.get("rate_tables", []))
    n_risks = len(analysis.get("risks", []))

    doc.add_paragraph(
        f"This document presents the design analysis for {plan_name}. "
        f"The AI-driven analysis identified {n_plans} compensation plan(s), "
        f"{n_components} plan component(s), {n_measures} performance measure(s), "
        f"{n_expressions} expression(s), and {n_rate_tables} rate table(s). "
        f"A total of {n_risks} risk(s) were flagged for review."
    )

    # Summary table
    _add_table(doc, ["Metric", "Count"], [
        ["Compensation Plans", str(n_plans)],
        ["Plan Components", str(n_components)],
        ["Performance Measures", str(n_measures)],
        ["Expressions", str(n_expressions)],
        ["Rate Tables", str(n_rate_tables)],
        ["Credit Categories", str(len(om.get("credit_categories", [])))],
        ["Risks Identified", str(n_risks)],
        ["Stakeholder Questions", str(len(analysis.get("stakeholder_questions", [])))],
    ])
    doc.add_paragraph("")

    # ── 2. Plan Structure & Mechanics ──
    doc.add_heading("2. Plan Structure & Mechanics", level=1)
    plan_structure = analysis.get("plan_structure", [])
    if plan_structure:
        for item in plan_structure:
            doc.add_heading(item.get("name", "Section"), level=2)
            doc.add_paragraph(item.get("description", ""))
            excerpt = item.get("raw_excerpt", "")
            if excerpt:
                p = doc.add_paragraph()
                run = p.add_run(f'Source excerpt: "{excerpt}"')
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(100, 116, 139)
    else:
        doc.add_paragraph("No plan structure elements were extracted from the document.")

    # Plan details from oracle_mapping
    plans = om.get("compensation_plans", [])
    if plans:
        doc.add_heading("Compensation Plan Details", level=2)
        _add_table(doc, ["Plan Name", "Start Date", "End Date", "Target Incentive", "Status"], [
            [
                p.get("name", ""),
                p.get("start_date", ""),
                p.get("end_date", ""),
                f"${p.get('target_incentive', 0):,.0f}",
                p.get("status", "Active"),
            ] for p in plans
        ])
        doc.add_paragraph("")

    # Component summary
    components = om.get("plan_components", [])
    if components:
        doc.add_heading("Plan Component Summary", level=2)
        _add_table(doc,
            ["Component", "Type", "Measure", "Calc Method", "Payout"],
            [
                [
                    c.get("plan_component_name", ""),
                    c.get("incentive_type", ""),
                    c.get("performance_measure_name", ""),
                    c.get("calculate_incentive", "Per interval"),
                    c.get("payout_frequency", "Period"),
                ] for c in components
            ],
        )
        doc.add_paragraph("")

    # ── 3. Risk Assessment ──
    doc.add_heading("3. Risk Assessment", level=1)
    risks = analysis.get("risks", [])
    if risks:
        # Summary counts
        high = sum(1 for r in risks if _severity_label(r.get("severity", "")) == "HIGH")
        med = sum(1 for r in risks if _severity_label(r.get("severity", "")) == "MEDIUM")
        low = sum(1 for r in risks if _severity_label(r.get("severity", "")) == "LOW")
        doc.add_paragraph(f"Risk summary: {high} high, {med} medium, {low} low severity items.")

        _add_table(doc, ["#", "Risk Title", "Severity", "Detail"], [
            [
                str(i + 1),
                r.get("title", ""),
                _severity_label(r.get("severity", "")),
                r.get("detail", ""),
            ] for i, r in enumerate(risks)
        ])
    else:
        doc.add_paragraph("No risks were identified in this analysis.")
    doc.add_paragraph("")

    # ── 4. Stakeholder Questions ──
    doc.add_heading("4. Stakeholder Questions", level=1)
    questions = analysis.get("stakeholder_questions", [])
    if questions:
        doc.add_paragraph(
            "The following questions should be resolved with stakeholders "
            "before proceeding to configuration:"
        )
        for i, q in enumerate(questions, 1):
            p = doc.add_paragraph(f"{i}. {q}")
            p.paragraph_format.space_after = Pt(4)
    else:
        doc.add_paragraph("No stakeholder questions were identified.")

    # ── 5. Oracle ICM Object Summary ──
    doc.add_heading("5. Oracle ICM Object Summary", level=1)
    doc.add_paragraph(
        "The following Oracle ICM objects were mapped from the plan document. "
        "These form the foundation for the Configuration Document and deployment."
    )

    # Expressions by category
    expressions = om.get("expressions", [])
    if expressions:
        doc.add_heading("Expression Formulas", level=2)
        _add_table(doc, ["Expression Name", "Category", "Description"], [
            [
                e.get("expression_name", ""),
                e.get("expression_category", "Earnings"),
                e.get("description", ""),
            ] for e in expressions
        ])
        doc.add_paragraph("")

    # Rate tables
    rate_tables = om.get("rate_tables", [])
    rate_table_rates = om.get("rate_table_rates", [])
    if rate_tables:
        doc.add_heading("Rate Tables", level=2)
        for rt in rate_tables:
            rt_name = rt.get("rate_table_name", "")
            doc.add_heading(rt_name, level=3)
            rates = [r for r in rate_table_rates if r.get("rate_table_name") == rt_name]
            if rates:
                _add_table(doc, ["Tier", "Min Amount", "Max Amount", "Rate"], [
                    [
                        str(r.get("tier_sequence", "")),
                        f"${r.get('minimum_amount', 0):,.0f}",
                        f"${r.get('maximum_amount', 0):,.0f}",
                        f"{r.get('rate_value', 0):.2%}" if r.get("rate_value", 0) < 1 else str(r.get("rate_value", "")),
                    ] for r in sorted(rates, key=lambda x: x.get("tier_sequence", 0))
                ])
                doc.add_paragraph("")

    # Performance measures
    measures = om.get("performance_measures", [])
    if measures:
        doc.add_heading("Performance Measures", level=2)
        _add_table(doc, ["Measure", "Unit", "Interval", "Formula Expression", "Credit Category"], [
            [
                m.get("name", ""),
                m.get("unit_of_measure", "AMOUNT"),
                m.get("performance_interval", "Quarterly"),
                m.get("measure_formula_expression_name", ""),
                m.get("credit_category_name", ""),
            ] for m in measures
        ])
        doc.add_paragraph("")

    # ── 6. Recommendations ──
    doc.add_heading("6. Recommendations", level=1)
    recs = analysis.get("reports_recommendations", [])
    if recs:
        for rec in recs:
            doc.add_paragraph(f"- {rec}", style="List Bullet")
    else:
        doc.add_paragraph("No specific recommendations were generated.")

    flex = analysis.get("operational_flexibility", [])
    if flex:
        doc.add_heading("Operational Flexibility", level=2)
        for f in flex:
            doc.add_paragraph(f"- {f}", style="List Bullet")

    # ── 7. Data Integration Requirements ──
    doc.add_heading("7. Data Integration Requirements", level=1)
    integrations = analysis.get("data_integrations", [])
    if integrations:
        for item in integrations:
            doc.add_paragraph(f"- {item}", style="List Bullet")
    else:
        doc.add_paragraph("No data integration requirements were specified.")

    # ── 8. Governance & Controls ──
    doc.add_heading("8. Governance & Controls", level=1)
    controls = analysis.get("governance_controls", [])
    if controls:
        for ctrl in controls:
            doc.add_paragraph(f"- {ctrl}", style="List Bullet")
    else:
        doc.add_paragraph("No governance controls were specified.")

    # Footer note
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Generated by ICM PlanLytics AI Plan Analytics")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(148, 163, 184)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ─── Configuration Document ─────────────────────────────────────────

def generate_config_document(analysis: dict, analysis_id: str) -> bytes:
    """
    Generate a Configuration Document (.docx) from analysis results.

    This is a detailed technical document containing all Oracle ICM objects
    with exact field values needed for deployment.

    Sections:
      1. Configuration Summary
      2. Deployment Sequence
      3. Credit Categories
      4. Rate Dimensions & Rate Tables
      5. Expressions (with formulas)
      6. Performance Measures & Goals
      7. Plan Components & Calculation Settings
      8. Compensation Plans
      9. Scorecards
     10. Validation Warnings
     11. REST API Deployment Reference
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    plan_name = _get_plan_name(analysis)
    _add_title_page(doc, "Configuration Document", plan_name, analysis_id)

    om = analysis.get("oracle_mapping", {})

    # ── TOC ──
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Configuration Summary",
        "2. Deployment Sequence",
        "3. Credit Categories",
        "4. Rate Dimensions & Rate Tables",
        "5. Expressions",
        "6. Performance Measures & Goals",
        "7. Plan Components & Calculation Settings",
        "8. Compensation Plans",
        "9. Scorecards",
        "10. Validation Warnings",
        "11. REST API Deployment Reference",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # ── 1. Configuration Summary ──
    doc.add_heading("1. Configuration Summary", level=1)
    summary_rows = [
        ["Object Type", "Count"],
    ]
    object_counts = [
        ("Credit Categories", len(om.get("credit_categories", []))),
        ("Rate Dimensions", len(om.get("rate_dimensions", []))),
        ("Rate Tables", len(om.get("rate_tables", []))),
        ("Rate Table Rates", len(om.get("rate_table_rates", []))),
        ("Expressions", len(om.get("expressions", []))),
        ("Performance Measures", len(om.get("performance_measures", []))),
        ("Performance Goals", len(om.get("performance_goals", []))),
        ("Plan Components", len(om.get("plan_components", []))),
        ("Compensation Plans", len(om.get("compensation_plans", []))),
        ("Scorecards", len(om.get("scorecards", []))),
        ("Calculation Settings", len(om.get("calculation_settings", []))),
    ]
    total = sum(c for _, c in object_counts)
    _add_table(doc, ["Object Type", "Count"],
        [[name, str(count)] for name, count in object_counts] +
        [["TOTAL", str(total)]]
    )
    doc.add_paragraph("")

    # ── 2. Deployment Sequence ──
    doc.add_heading("2. Deployment Sequence", level=1)
    doc.add_paragraph(
        "Oracle ICM objects must be created in a specific order to satisfy "
        "dependency requirements. The following is the required deployment sequence:"
    )
    deploy_steps = [
        ("Step 1", "Credit Categories", "POST /incentiveCompensationCreditCategories", "Define credit categories for transaction classification"),
        ("Step 2", "Rate Dimensions", "POST /rateDimensions", "Create tier boundaries for rate table lookups"),
        ("Step 3", "Rate Tables", "POST /rateTables", "Create rate table containers and associate dimensions"),
        ("Step 4", "Expressions", "POST /incentiveCompensationExpressions", "Create attainment, earnings, and rate dimension input formulas"),
        ("Step 5", "Performance Measures", "POST /incentiveCompensationPerformanceMeasures", "Define measures with formula expressions and credit categories"),
        ("Step 6", "Plan Components", "POST /planComponents", "Link measures, rate tables, and formulas to components"),
        ("Step 7", "Compensation Plans", "POST /compensationPlans", "Create top-level plan containers"),
        ("Step 8", "Attach Components", "POST /compensationPlans/{id}/child/CompensationPlanComponents", "Associate plan components with compensation plans"),
    ]
    _add_table(doc, ["Step", "Object", "REST Endpoint", "Description"],
        [[s[0], s[1], s[2], s[3]] for s in deploy_steps]
    )
    doc.add_paragraph("")

    # ── 3. Credit Categories ──
    doc.add_heading("3. Credit Categories", level=1)
    ccs = om.get("credit_categories", [])
    if ccs:
        _add_table(doc, ["Credit Category Name", "Description", "Org ID"], [
            [
                cc.get("credit_category_name", ""),
                cc.get("description", ""),
                str(cc.get("org_id", 300000046987012)),
            ] for cc in ccs
        ])
    else:
        doc.add_paragraph("No credit categories defined.")
    doc.add_paragraph("")

    # ── 4. Rate Dimensions & Rate Tables ──
    doc.add_heading("4. Rate Dimensions & Rate Tables", level=1)

    rds = om.get("rate_dimensions", [])
    if rds:
        doc.add_heading("Rate Dimensions", level=2)
        _add_table(doc,
            ["Dimension Name", "Type", "Tier", "Min Amount", "Max Amount", "Org ID"],
            [
                [
                    rd.get("rate_dimension_name", ""),
                    rd.get("rate_dimension_type", "AMOUNT"),
                    str(rd.get("tier_sequence", 1)),
                    str(rd.get("minimum_amount", 0)),
                    str(rd.get("maximum_amount", 999999)),
                    str(rd.get("org_id", 300000046987012)),
                ] for rd in rds
            ],
        )
        doc.add_paragraph("")

    rts = om.get("rate_tables", [])
    if rts:
        doc.add_heading("Rate Tables", level=2)
        _add_table(doc, ["Rate Table Name", "Type", "Display Name", "Org ID"], [
            [
                rt.get("rate_table_name", ""),
                rt.get("rate_table_type", "Sales"),
                rt.get("display_name", ""),
                str(rt.get("org_id", 300000046987012)),
            ] for rt in rts
        ])
        doc.add_paragraph("")

    rtrs = om.get("rate_table_rates", [])
    if rtrs:
        doc.add_heading("Rate Table Rates", level=2)
        _add_table(doc, ["Rate Table Name", "Tier", "Min Amount", "Max Amount", "Rate Value"], [
            [
                r.get("rate_table_name", ""),
                str(r.get("tier_sequence", 1)),
                str(r.get("minimum_amount", 0)),
                str(r.get("maximum_amount", 999999)),
                str(r.get("rate_value", 0)),
            ] for r in rtrs
        ])
    doc.add_paragraph("")

    # ── 5. Expressions ──
    doc.add_heading("5. Expressions", level=1)
    doc.add_paragraph(
        "Expressions define the calculation logic for attainment, earnings, "
        "rate dimension input, and weighted score computations."
    )
    exprs = om.get("expressions", [])
    if exprs:
        # Group by category
        for category in ("Attainment", "Earnings", "RateDimensionInput", "Weighted"):
            cat_exprs = [e for e in exprs if e.get("expression_category") == category]
            if not cat_exprs:
                continue
            doc.add_heading(f"{category} Expressions", level=2)
            for expr in cat_exprs:
                doc.add_heading(expr.get("expression_name", ""), level=3)
                _add_table(doc, ["Field", "Value"], [
                    ["Expression Name", expr.get("expression_name", "")],
                    ["Category", expr.get("expression_category", "")],
                    ["Type", expr.get("expression_type", "Calculation")],
                    ["Description / Formula", expr.get("description", "")],
                    ["Measure Name", str(expr.get("measure_name", ""))],
                    ["Basic Attributes Group", str(expr.get("basic_attributes_group", ""))],
                    ["Basic Attribute Name", str(expr.get("basic_attribute_name", ""))],
                    ["Measure Result Attribute", str(expr.get("measure_result_attribute", ""))],
                    ["Plan Component Name", str(expr.get("plan_component_name", ""))],
                    ["Expression Operator", str(expr.get("expression_operator", ""))],
                    ["Constant Value", str(expr.get("constant_value", ""))],
                    ["Sequence", str(expr.get("sequence", ""))],
                ])
                doc.add_paragraph("")
    else:
        doc.add_paragraph("No expressions defined.")

    # ── 6. Performance Measures & Goals ──
    doc.add_heading("6. Performance Measures & Goals", level=1)
    pms = om.get("performance_measures", [])
    if pms:
        doc.add_heading("Performance Measures", level=2)
        for pm in pms:
            doc.add_heading(pm.get("name", ""), level=3)
            _add_table(doc, ["Field", "Value"], [
                ["Name", pm.get("name", "")],
                ["Description", pm.get("description", "")],
                ["Unit of Measure", pm.get("unit_of_measure", "AMOUNT")],
                ["Performance Interval", pm.get("performance_interval", "Quarterly")],
                ["Formula Expression", str(pm.get("measure_formula_expression_name", ""))],
                ["Credit Category", pm.get("credit_category_name", "")],
                ["Process Transactions", pm.get("process_transactions", "Yes")],
                ["Running Total", pm.get("running_total_flag", "N")],
                ["Active", pm.get("active_flag", "Y")],
                ["Fiscal Year", str(pm.get("f_year", ""))],
                ["Start Date", pm.get("start_date", "")],
                ["End Date", pm.get("end_date", "")],
                ["Scorecard Rate Table", str(pm.get("scorecard_rate_table_name", ""))],
                ["Org ID", str(pm.get("org_id", 300000046987012))],
            ])
            doc.add_paragraph("")

    goals = om.get("performance_goals", [])
    if goals:
        doc.add_heading("Performance Goals", level=2)
        _add_table(doc, ["Measure Name", "Goal Interval", "Goal Target"], [
            [
                g.get("performance_measure_name", ""),
                g.get("goal_interval", "Quarterly"),
                f"${g.get('goal_target', 0):,.0f}" if isinstance(g.get("goal_target"), (int, float)) else str(g.get("goal_target", "")),
            ] for g in goals
        ])
        doc.add_paragraph("")

    # ── 7. Plan Components & Calculation Settings ──
    doc.add_heading("7. Plan Components & Calculation Settings", level=1)
    pcs = om.get("plan_components", [])
    calc_settings = om.get("calculation_settings", [])
    # Index calc settings by component name
    cs_map = {cs.get("plan_component_name", ""): cs for cs in calc_settings}

    if pcs:
        for pc in pcs:
            pc_name = pc.get("plan_component_name", "")
            doc.add_heading(pc_name, level=2)

            # Component details
            _add_table(doc, ["Field", "Value"], [
                ["Plan Name", pc.get("plan_name", "")],
                ["Component Name", pc_name],
                ["Incentive Type", pc.get("incentive_type", "Sales")],
                ["Calculation Method", pc.get("calculation_method", "Tiered")],
                ["Performance Measure", pc.get("performance_measure_name", "")],
                ["Measure Weight", str(pc.get("performance_measure_weight", 1.0))],
                ["Rate Table", pc.get("rate_table_name", "")],
                ["Incentive Formula", pc.get("incentive_formula_expression", "")],
                ["Earning Basis", pc.get("earning_basis", "Amount")],
                ["Calculation Sequence", str(pc.get("calculation_sequence", 1))],
                ["Start Date", pc.get("start_date", "")],
                ["End Date", pc.get("end_date", "")],
                ["Org ID", str(pc.get("org_id", 300000046987012))],
            ])
            doc.add_paragraph("")

            # Matching calculation settings
            cs = cs_map.get(pc_name)
            if cs:
                doc.add_heading(f"Calculation Settings — {pc_name}", level=3)
                _add_table(doc, ["Setting", "Value"], [
                    ["Calculate Incentive", cs.get("calculate_incentive", "Per interval")],
                    ["Process Transactions", cs.get("process_transactions", "Grouped by interval")],
                    ["Payout Frequency", cs.get("payout_frequency", "Period")],
                    ["Split Attainment", cs.get("split_attainment", "No")],
                    ["Fixed Within Tier", cs.get("fixed_within_tier", "No")],
                    ["True Up", cs.get("true_up", "No")],
                    ["True Up Reset Interval", str(cs.get("true_up_reset_interval", "N/A"))],
                    ["Include Indirect Credits", cs.get("include_indirect_credits", "None")],
                    ["Running Total", cs.get("running_total", "No")],
                ])
                doc.add_paragraph("")

            # Rate dimension input expression
            rdi = pc.get("rate_dimension_input_expression")
            if rdi:
                p = doc.add_paragraph()
                run = p.add_run(f"Rate Dimension Input Expression: {rdi}")
                run.italic = True
                doc.add_paragraph("")
    else:
        doc.add_paragraph("No plan components defined.")

    # ── 8. Compensation Plans ──
    doc.add_heading("8. Compensation Plans", level=1)
    plans = om.get("compensation_plans", [])
    if plans:
        for plan in plans:
            doc.add_heading(plan.get("name", ""), level=2)
            _add_table(doc, ["Field", "Value"], [
                ["Name", plan.get("name", "")],
                ["Display Name", plan.get("display_name", "")],
                ["Description", plan.get("description", "")],
                ["Start Date", plan.get("start_date", "")],
                ["End Date", plan.get("end_date", "")],
                ["Status", plan.get("status", "Active")],
                ["Target Incentive", f"${plan.get('target_incentive', 0):,.0f}" if isinstance(plan.get("target_incentive"), (int, float)) else str(plan.get("target_incentive", ""))],
                ["Org ID", str(plan.get("org_id", 300000046987012))],
            ])
            doc.add_paragraph("")

            # List attached components
            attached = [pc for pc in pcs if pc.get("plan_name") == plan.get("name")]
            if attached:
                doc.add_paragraph("Attached Plan Components:")
                for apc in attached:
                    doc.add_paragraph(
                        f"- {apc.get('plan_component_name', '')} "
                        f"(Seq: {apc.get('calculation_sequence', 1)}, "
                        f"Measure: {apc.get('performance_measure_name', '')})",
                        style="List Bullet",
                    )
                doc.add_paragraph("")
    else:
        doc.add_paragraph("No compensation plans defined.")

    # ── 9. Scorecards ──
    doc.add_heading("9. Scorecards", level=1)
    scorecards = om.get("scorecards", [])
    if scorecards:
        _add_table(doc,
            ["Scorecard Name", "Performance Measure", "Rate Table", "Input Expression", "Description"],
            [
                [
                    sc.get("scorecard_name", ""),
                    sc.get("performance_measure_name", ""),
                    sc.get("rate_table_name", ""),
                    sc.get("input_expression_name", ""),
                    sc.get("description", ""),
                ] for sc in scorecards
            ],
        )
    else:
        doc.add_paragraph("No scorecards defined for this plan.")
    doc.add_paragraph("")

    # ── 10. Validation Warnings ──
    doc.add_heading("10. Validation Warnings", level=1)
    warnings = analysis.get("validation_warnings", [])
    if warnings:
        doc.add_paragraph(
            "The following cross-reference validation issues were detected. "
            "These should be resolved before deployment:"
        )
        for w in warnings:
            doc.add_paragraph(f"- {w}", style="List Bullet")
    else:
        p = doc.add_paragraph("No validation warnings. All cross-references are consistent.")
        run = p.runs[0]
        run.font.color.rgb = RGBColor(34, 197, 94)
    doc.add_paragraph("")

    # ── 11. REST API Deployment Reference ──
    doc.add_heading("11. REST API Deployment Reference", level=1)
    doc.add_paragraph(
        "The following REST API endpoints are used for automated deployment "
        "to Oracle Fusion Incentive Compensation Management."
    )
    doc.add_paragraph(
        "Base URL: https://{instance}.fa.{region}.oraclecloud.com"
        "/fscmRestApi/resources/11.13.18.05/"
    )
    _add_table(doc, ["Object", "Method", "Endpoint", "Key Fields"], [
        ["Compensation Plans", "POST", "/compensationPlans", "Name, StartDate, EndDate, OrgId, TargetIncentive"],
        ["Expressions", "POST", "/incentiveCompensationExpressions", "Name, OrgId, ExpressionDetails[]"],
        ["Plan Components", "POST", "/planComponents", "Name, measures, formulas, rate tables (children)"],
        ["Performance Measures", "GET", "/incentiveCompensationPerformanceMeasures", "Name, UnitOfMeasure, formula ref"],
        ["Rate Tables", "GET/POST", "/rateTables", "Name, rateTableDimensions (children)"],
        ["Plan Assignments", "POST", "/compensationPlans/{id}/child/CompensationPlanAssignments", "ParticipantId, AssignmentType"],
    ])
    doc.add_paragraph("")
    doc.add_paragraph(
        "Note: Expressions are created with status INVALID by default and "
        "must be validated after creation. EndDate cannot be set to null once "
        "assigned. Date format: yyyy-mm-dd."
    )

    # Footer
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Generated by ICM PlanLytics AI Plan Analytics")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(148, 163, 184)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ─── Utility ────────────────────────────────────────────────────────

def _get_plan_name(analysis: dict) -> str:
    """Extract a readable plan name from analysis data."""
    om = analysis.get("oracle_mapping", {})
    plans = om.get("compensation_plans", [])
    if plans:
        return plans[0].get("name", "Compensation Plan")
    return "Compensation Plan Analysis"
